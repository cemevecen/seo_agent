"""Ayılma hemşireleri — aylık çalışma çizelgesi motoru.

Hücre kodları (örnek SSE ile uyumlu):
  8  → 08:00–16:00 (8 saat)
  16 → 16:00–08:00 (16 saat)
  24 → 08:00–08:00 (24 saat)
  Yİ → yıllık izin
  RP → rapor
  İST → özel gün isteği / rezervasyon (çalıştırılmaz)
  '' → boş / dinlenme
"""

from __future__ import annotations

import calendar
import random
from dataclasses import dataclass
from datetime import date
from typing import Any

LEAD_NURSE = "Gülten Çelik"
JIN_NURSE_LABEL = "JIN"

STAFF_NURSES: tuple[str, ...] = (
    "Nuray Durna",
    "Sema Evecen",
    "Şengül Zamur",
    "Emine Türker",
    "Semanur Çınar",
    "Rabia Kumtepe",
)

ALL_NURSES: tuple[str, ...] = (LEAD_NURSE, *STAFF_NURSES)

LEAVE_CODES = frozenset({"Yİ", "RP", "İST"})
WORK_CODES = frozenset({"8", "16", "24"})
JIN_WORK_CODES = frozenset({"16", "24"})
MAX_MONTHLY_HOURS = 300
# Yıllık izin günü = 8 saat mesai kullanılmış sayılır
YI_DAY_HOURS = 8
# Personel arası toplam (mesai+Yİ) / fazla mesai bandı hedefi (~16s)
HOURS_BALANCE_TOLERANCE = 16
# Kişi başı düz 8 sayısı (3×8≈24s + boş gün üretir; az tut)
EIGHT_PER_PERSON_MIN = 2
EIGHT_PER_PERSON_TARGET = 3
EIGHT_PER_PERSON_MAX = 4
# Gün aşırı zinciri: 24+boş+24+… en fazla 3×24 (4. yasak; izin yoğun olsa da)
GUN_ASIRI_STREAK_MAX = 3
# Üst üste «8»: tercih yok; mecbur kalınırsa en fazla 2 gün
CONSECUTIVE_8_STREAK_MAX = 2
# Haftada bu kadar Yİ/RP/İST hücresi varsa soft cezalar gevşer (streak tavanı gevşemez)
LEAVE_HEAVY_WEEK_THRESHOLD = 6


def _yi_hours_from_grid(grid: dict[str, dict[str, str]], name: str, days: list[DayMeta]) -> int:
    return sum(YI_DAY_HOURS for dm in days if grid[name].get(dm.iso, "") == "Yİ")


def _shift_hours_from_grid(grid: dict[str, dict[str, str]], name: str, days: list[DayMeta]) -> int:
    return sum(_hours_for(grid[name].get(dm.iso, "")) for dm in days)


@dataclass(frozen=True)
class DayMeta:
    day: int
    iso: str
    weekday: int  # Mon=0 … Sun=6
    is_weekend: bool
    is_weekday: bool


def _hours_for(code: str | None) -> int:
    c = (code or "").strip().upper()
    if c == "8":
        return 8
    if c == "16":
        return 16
    if c == "24":
        return 24
    return 0


def _norm_code(raw: Any) -> str:
    if raw is None:
        return ""
    s = str(raw).strip()
    if not s:
        return ""
    up = s.upper()
    if up in ("YI", "Yİ"):
        return "Yİ"
    if up == "RP":
        return "RP"
    if up in ("IST", "İST"):
        return "İST"
    if s in ("8", "16", "24"):
        return s
    return ""


def month_days(year: int, month: int) -> list[DayMeta]:
    n = calendar.monthrange(year, month)[1]
    out: list[DayMeta] = []
    for d in range(1, n + 1):
        dt = date(year, month, d)
        wd = dt.weekday()
        out.append(
            DayMeta(
                day=d,
                iso=dt.isoformat(),
                weekday=wd,
                is_weekend=wd >= 5,
                is_weekday=wd < 5,
            )
        )
    return out


def count_weekdays(year: int, month: int) -> int:
    return sum(1 for dm in month_days(year, month) if dm.is_weekday)


def ideal_hours(year: int, month: int) -> int:
    return count_weekdays(year, month) * 8


def _empty_grid(year: int, month: int) -> dict[str, dict[str, str]]:
    days = month_days(year, month)
    grid: dict[str, dict[str, str]] = {}
    for name in ALL_NURSES:
        grid[name] = {dm.iso: "" for dm in days}
    return grid


def _apply_leaves(
    grid: dict[str, dict[str, str]],
    leaves: dict[str, dict[str, str]] | None,
) -> None:
    if not leaves:
        return
    for name, by_day in leaves.items():
        if name not in grid or not isinstance(by_day, dict):
            continue
        for iso, code in by_day.items():
            nc = _norm_code(code)
            if nc in LEAVE_CODES and iso in grid[name]:
                grid[name][iso] = nc


def _blocked_by_rest(
    name: str,
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
    *,
    prefer_48h_after_24: bool,
) -> bool:
    """Önceki gece 16/24 bitişi bugünün sabahına denk gelir → bugün başlama."""
    if day_index <= 0:
        return False
    prev = days[day_index - 1]
    prev_code = grid[name].get(prev.iso, "")
    if prev_code == "24":
        return True
    if prev_code == "16":
        return True
    return False


def _next_day_blocks_24(
    name: str,
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
) -> bool:
    """Yarın zaten 16/24 varsa bugün 24 yazmak arka arkaya nöbet üretir."""
    if day_index + 1 >= len(days):
        return False
    return grid[name].get(days[day_index + 1].iso, "") in ("16", "24")


def _cannot_assign_24(
    name: str,
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
    *,
    prefer_48h_after_24: bool,
) -> bool:
    """Arka arkaya 24/16 engeli (önceki veya sonraki gün)."""
    return _blocked_by_rest(
        name, day_index, days, grid, prefer_48h_after_24=prefer_48h_after_24
    ) or _next_day_blocks_24(name, day_index, days, grid)


def _rest_penalty(
    name: str,
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
) -> int:
    """Düşük = daha iyi aday. 24 sonrası 2. gün hâlâ hafif cezalı (ideal 48s)."""
    if day_index >= 1 and grid[name].get(days[day_index - 1].iso, "") in ("16", "24"):
        return 999
    if day_index >= 2 and grid[name].get(days[day_index - 2].iso, "") == "24":
        return 25
    return 0


def _gun_asiri_24_penalty(
    name: str,
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
) -> int:
    """24 + (1 gün boş) + 24 = 'gün aşırı nöbet' — yumuşak ceza (16 yazmak için değil)."""
    if day_index < 2:
        return 0
    if grid[name].get(days[day_index - 2].iso, "") != "24":
        return 0
    mid = grid[name].get(days[day_index - 1].iso, "")
    if mid in ("", "Yİ", "RP", "İST"):
        return 1
    return 0


def _gun_asiri_streak_if_24(
    name: str,
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
) -> int:
    """Bugün 24 yazılırsa gün aşırı zincirinde toplam kaç 24 olur (geri+ileri)."""
    gap = ("", "Yİ", "RP", "İST")

    back = 1  # today
    j = day_index
    while j >= 2:
        if grid[name].get(days[j - 2].iso, "") != "24":
            break
        mid = grid[name].get(days[j - 1].iso, "")
        if mid not in gap:
            break
        back += 1
        j -= 2

    forward = 0
    j = day_index
    while j + 2 < len(days):
        mid = grid[name].get(days[j + 1].iso, "")
        if mid not in gap:
            break
        if grid[name].get(days[j + 2].iso, "") != "24":
            break
        forward += 1
        j += 2

    return back + forward


def _week_leave_cell_count(
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
) -> int:
    """İçinde bulunulan Pzt–Paz haftasındaki personel izin/istek hücresi sayısı."""
    wd = days[day_index].weekday
    start = max(0, day_index - wd)
    end = min(len(days) - 1, start + 6)
    n = 0
    for i in range(start, end + 1):
        iso = days[i].iso
        for name in STAFF_NURSES:
            if grid[name].get(iso, "") in ("Yİ", "RP", "İST"):
                n += 1
    return n


def _week_is_leave_heavy(
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
) -> bool:
    return _week_leave_cell_count(day_index, days, grid) >= LEAVE_HEAVY_WEEK_THRESHOLD


def _prefer_8_after_24_gap(
    name: str,
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
) -> int:
    """24 sonrası boş günde 8 serpiştir (düşük = tercih)."""
    if day_index >= 2 and grid[name].get(days[day_index - 2].iso, "") == "24":
        mid = grid[name].get(days[day_index - 1].iso, "")
        if mid in ("", "Yİ", "RP", "İST"):
            return 0
    if day_index >= 3 and grid[name].get(days[day_index - 3].iso, "") == "24":
        return 1
    return 4


def _prev_day_is_8(
    name: str,
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
) -> bool:
    if day_index <= 0:
        return False
    return grid[name].get(days[day_index - 1].iso, "") == "8"


def _consecutive_8_streak_if_8(
    name: str,
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
) -> int:
    """Bugün «8» yazılırsa kaç gün üst üste 8 olur."""
    streak = 1
    j = day_index - 1
    while j >= 0 and grid[name].get(days[j].iso, "") == "8":
        streak += 1
        j -= 1
    return streak


def _can_assign_8(
    name: str,
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
    *,
    relax_consecutive: bool = False,
) -> bool:
    """Üst üste 8 sınırı. Özel koşul (pin) 8'i olan kişide art arda serbest."""
    if relax_consecutive:
        return True
    return (
        _consecutive_8_streak_if_8(name, day_index, days, grid) <= CONSECUTIVE_8_STREAK_MAX
    )


def _max_consecutive_8_streak(
    name: str,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
) -> int:
    best = 0
    streak = 0
    for dm in days:
        if grid[name].get(dm.iso, "") == "8":
            streak += 1
            best = max(best, streak)
        else:
            streak = 0
    return best


def _first_day_after_leave(
    name: str,
    day_index: int,
    days: list[DayMeta],
    grid: dict[str, dict[str, str]],
) -> bool:
    """Dün Yİ/RP (izin/rapor); bugün mesai günü — dönüşte nöbet tercih."""
    if day_index <= 0:
        return False
    prev = grid[name].get(days[day_index - 1].iso, "")
    if prev not in ("Yİ", "RP"):
        return False
    today = grid[name].get(days[day_index].iso, "")
    return today not in LEAVE_CODES


def _count_consecutive_8_runs(
    grid: dict[str, dict[str, str]],
    days: list[DayMeta],
    *,
    staff_only: bool = True,
) -> int:
    """Ardışık 8 çifti sayısı (üst üste 8+8)."""
    names = list(STAFF_NURSES) if staff_only else list(ALL_NURSES)
    runs = 0
    for name in names:
        for i in range(1, len(days)):
            if grid[name].get(days[i].iso, "") != "8":
                continue
            if grid[name].get(days[i - 1].iso, "") == "8":
                runs += 1
    return runs


NIGHT_SHIFTS_PER_DAY = 2


def _staff_night_count(grid: dict[str, dict[str, str]], iso: str) -> int:
    return sum(1 for n in STAFF_NURSES if grid[n].get(iso, "") in ("16", "24"))


def _jin_night_slots(slots: list[str]) -> int:
    return sum(1 for c in slots if c in JIN_WORK_CODES)


def _format_jin_cell(slots: list[str]) -> str:
    if not slots:
        return ""
    counts: dict[str, int] = {}
    for code in slots:
        if code not in JIN_WORK_CODES:
            continue
        counts[code] = counts.get(code, 0) + 1
    parts: list[str] = []
    for code in ("16", "24"):
        n = counts.get(code, 0)
        if n == 1:
            parts.append(code)
        elif n > 1:
            parts.append(f"{code}×{n}")
    return "+".join(parts)


def resolve_special_day_sets(
    year: int,
    month: int,
    special_rules: list[dict[str, Any]] | None,
) -> tuple[dict[str, set[str]], dict[str, set[str]]]:
    """Özel koşul → çalışsın / çalışmasın ISO kümeleri."""
    work: dict[str, set[str]] = {n: set() for n in STAFF_NURSES}
    avoid: dict[str, set[str]] = {n: set() for n in STAFF_NURSES}
    if not special_rules:
        return work, avoid
    days = month_days(year, month)
    by_weekday: dict[int, list[str]] = {}
    for dm in days:
        by_weekday.setdefault(dm.weekday, []).append(dm.iso)
    iso_to_wd = {dm.iso: dm.weekday for dm in days}
    month_isos = set(iso_to_wd)

    for raw in special_rules:
        if not isinstance(raw, dict):
            continue
        name = str(raw.get("name") or "").strip()
        if name not in STAFF_NURSES:
            continue
        mode = str(raw.get("mode") or "").strip().lower()
        if mode in ("calissin", "çalışsın", "work"):
            mode = "work"
        elif mode in ("calismasin", "çalışmasın", "avoid"):
            mode = "avoid"
        else:
            continue
        dates_raw = raw.get("dates") or []
        if not isinstance(dates_raw, list):
            continue
        picked = [str(x).strip() for x in dates_raw if str(x).strip()]
        weekly = bool(raw.get("weekly"))
        expanded: set[str] = set()
        for iso in picked:
            if weekly and iso in iso_to_wd:
                expanded.update(by_weekday.get(iso_to_wd[iso], []))
            elif iso in month_isos:
                expanded.add(iso)
        if mode == "work":
            work[name] |= expanded
        else:
            avoid[name] |= expanded
    for n in STAFF_NURSES:
        clash = work[n] & avoid[n]
        if clash:
            work[n] -= clash
    return work, avoid


def resolve_special_pins(
    year: int,
    month: int,
    special_rules: list[dict[str, Any]] | None,
) -> dict[str, dict[str, str]]:
    """Özel koşul → kişi/gün sabit vardiya (8/16/24)."""
    pins: dict[str, dict[str, str]] = {n: {} for n in STAFF_NURSES}
    if not special_rules:
        return pins
    days = month_days(year, month)
    by_weekday: dict[int, list[str]] = {}
    for dm in days:
        by_weekday.setdefault(dm.weekday, []).append(dm.iso)
    iso_to_wd = {dm.iso: dm.weekday for dm in days}
    month_isos = set(iso_to_wd)

    for raw in special_rules:
        if not isinstance(raw, dict):
            continue
        name = str(raw.get("name") or "").strip()
        if name not in STAFF_NURSES:
            continue
        mode = str(raw.get("mode") or "").strip().lower()
        if mode not in ("pin", "shift", "sabit", "vardiya", "fixed"):
            continue
        weekly = bool(raw.get("weekly"))
        shifts_raw = raw.get("shifts")
        pairs: list[tuple[str, str]] = []
        if isinstance(shifts_raw, dict):
            for iso, code in shifts_raw.items():
                nc = _norm_code(code)
                if nc in WORK_CODES and str(iso).strip() in month_isos:
                    pairs.append((str(iso).strip(), nc))
        else:
            code = _norm_code(raw.get("code"))
            if code not in WORK_CODES:
                continue
            dates_raw = raw.get("dates") or []
            if not isinstance(dates_raw, list):
                continue
            for iso in dates_raw:
                s = str(iso).strip()
                if s in month_isos:
                    pairs.append((s, code))
        for iso, code in pairs:
            targets = (
                by_weekday.get(iso_to_wd[iso], [])
                if weekly and iso in iso_to_wd
                else [iso]
            )
            for t in targets:
                pins[name][t] = code
    return pins


def resolve_jin_coverage(
    year: int,
    month: int,
    special_rules: list[dict[str, Any]] | None,
) -> dict[str, list[str]]:
    """JIN (jinekoloji destek) → ISO → [16|24, …] slot listesi (8 yok)."""
    coverage: dict[str, list[str]] = {}
    if not special_rules:
        return coverage
    days = month_days(year, month)
    by_weekday: dict[int, list[str]] = {}
    for dm in days:
        by_weekday.setdefault(dm.weekday, []).append(dm.iso)
    iso_to_wd = {dm.iso: dm.weekday for dm in days}
    month_isos = set(iso_to_wd)

    def _append_slots(iso: str, codes: list[str]) -> None:
        if not codes:
            return
        coverage.setdefault(iso, []).extend(codes)

    for raw in special_rules:
        if not isinstance(raw, dict):
            continue
        mode = str(raw.get("mode") or "").strip().lower()
        if mode not in ("jin", "jinekoloji", "destek", "jin_destek"):
            continue
        weekly = bool(raw.get("weekly"))
        slots_raw = raw.get("slots")
        pairs: list[tuple[str, list[str]]] = []
        if isinstance(slots_raw, dict):
            for iso, val in slots_raw.items():
                s_iso = str(iso).strip()
                if s_iso not in month_isos:
                    continue
                if isinstance(val, list):
                    codes = [_norm_code(c) for c in val]
                else:
                    codes = [_norm_code(val)]
                codes = [c for c in codes if c in JIN_WORK_CODES]
                if codes:
                    pairs.append((s_iso, codes))
        shifts_raw = raw.get("shifts")
        if isinstance(shifts_raw, dict):
            for iso, code in shifts_raw.items():
                s_iso = str(iso).strip()
                nc = _norm_code(code)
                if nc in JIN_WORK_CODES and s_iso in month_isos:
                    pairs.append((s_iso, [nc]))
        if not pairs:
            code = _norm_code(raw.get("code"))
            dates_raw = raw.get("dates") or []
            if code in JIN_WORK_CODES and isinstance(dates_raw, list):
                for iso in dates_raw:
                    s_iso = str(iso).strip()
                    if s_iso in month_isos:
                        pairs.append((s_iso, [code]))
        for iso, codes in pairs:
            targets = (
                by_weekday.get(iso_to_wd[iso], [])
                if weekly and iso in iso_to_wd
                else [iso]
            )
            for t in targets:
                _append_slots(t, codes)
    return coverage


def _apply_special_pins(
    grid: dict[str, dict[str, str]],
    pins: dict[str, dict[str, str]],
    hours: dict[str, int],
    n8: dict[str, int],
    n16: dict[str, int],
    n24: dict[str, int],
) -> None:
    """Sabit 8/16/24 yaz — izin dahil üzerine yazar (sorumlu pin öncelikli)."""
    for name, by_day in pins.items():
        for iso, code in by_day.items():
            cur = grid[name].get(iso, "")
            if cur == code:
                continue
            if cur in WORK_CODES:
                hours[name] -= _hours_for(cur)
                if cur == "8":
                    n8[name] -= 1
                elif cur == "16":
                    n16[name] -= 1
                elif cur == "24":
                    n24[name] -= 1
            grid[name][iso] = code
            hours[name] += _hours_for(code)
            if code == "8":
                n8[name] += 1
            elif code == "16":
                n16[name] += 1
            elif code == "24":
                n24[name] += 1


def generate_ayilma_schedule(
    year: int,
    month: int,
    *,
    leaves: dict[str, dict[str, str]] | None = None,
    day_only: list[str] | None = None,
    prefer_48h_after_24: bool = True,
    special_rules: list[dict[str, Any]] | None = None,
    variant: int = 0,
) -> dict[str, Any]:
    """6 personel + sorumlu.

    Hesap motoru: aa3ee655 (2026-08-26 21:40) tabanı + özel koşul.
    variant>0 → eşitlikte farklı aday (Yeniden oluştur).
    special_rules: çalışsın / çalışmasın / sabit vardiya (8/16/24) / JIN (16/24).

    Gülten panelde boş satır; hesaba karışmaz.
    Hafta içi: 1×«8» + 2×«24». Hafta sonu: yalnız 2×«24» (kat-1 / 8 yok).
    Yİ/RP bitişinin ertesi günü o kişiye nöbet (24) tercih.
    Düz «8» kişi başı aylık ~2–4 (hedef 3). Fazla mesai personelde aynı ~16s bantta.
    Üst üste «8» kaçınılır; mecbur kalınırsa en fazla 2 gün.
    Özel koşul ile pinlenen «8» birincil: o kişide art arda 8 serbest; motor ezmez.
    Hafta içi 8 pin / hafta içi 8 ile mesaisi dolan → hafta sonu nöbet (ek mesai) yok.
    Gün aşırı zinciri en fazla 3×24 (kati).
    Arka arkaya 24 yasak. «16» motor yazmaz — yalnız sorumlu pin ile.
    """
    if not (1 <= month <= 12):
        raise ValueError("month 1–12 olmalı")
    if year < 2000 or year > 2100:
        raise ValueError("year geçersiz")

    days = month_days(year, month)
    grid = _empty_grid(year, month)
    _apply_leaves(grid, leaves)
    day_only_set = {str(x).strip() for x in (day_only or []) if str(x).strip()}
    prefer_work, force_avoid = resolve_special_day_sets(year, month, special_rules)
    pins = resolve_special_pins(year, month, special_rules)
    jin_coverage = resolve_jin_coverage(year, month, special_rules)

    def _night_target(iso: str) -> int:
        return max(0, NIGHT_SHIFTS_PER_DAY - _jin_night_slots(jin_coverage.get(iso, [])))

    variant_i = max(0, int(variant or 0))
    rng = random.Random((year * 100 + month) * 10007 + variant_i * 7919 + 17)
    # variant=0: klasik isim sırası (değişmez); variant>0: kadro/gün sırası + eşitlik karışır
    staff_order = list(STAFF_NURSES)
    day_order = list(range(len(days)))
    if variant_i:
        rng.shuffle(staff_order)
        rng.shuffle(day_order)
    order_idx = {n: i for i, n in enumerate(staff_order)}
    person_tie = {n: (rng.random() if variant_i else 0.0) for n in STAFF_NURSES}

    def _tie(n: str) -> float:
        return person_tie[n]

    def _order(n: str) -> int:
        return order_idx[n] if variant_i else 0

    # Görsel: Gülten boş (hesaba dahil değil)
    for dm in days:
        grid[LEAD_NURSE][dm.iso] = ""

    hours = {n: 0 for n in STAFF_NURSES}
    n8 = {n: 0 for n in STAFF_NURSES}
    n24 = {n: 0 for n in STAFF_NURSES}
    n16 = {n: 0 for n in STAFF_NURSES}
    warnings: list[str] = []
    # Pin izin/İST üzerine yazar; Yİ saatleri pin sonrası hesaplanır
    _apply_special_pins(grid, pins, hours, n8, n16, n24)
    pinned_cells = {
        (name, iso) for name, by_day in pins.items() for iso in by_day
    }
    # Özel koşuldan 8 pinlenen kişi: art arda 8 birincil olarak serbest
    relax_consec_8 = {
        n for n, by_day in pins.items() if any(code == "8" for code in by_day.values())
    }

    def _ok8(n: str, day_index: int) -> bool:
        return _can_assign_8(
            n,
            day_index,
            days,
            grid,
            relax_consecutive=(n in relax_consec_8),
        )

    ideal = ideal_hours(year, month)
    yi_hours = {n: _yi_hours_from_grid(grid, n, days) for n in STAFF_NURSES}
    min_shift = {n: max(0, ideal - yi_hours[n]) for n in STAFF_NURSES}

    def accounted(n: str) -> int:
        """Mesai + Yİ (8s/gün) — dengelenecek toplam."""
        return hours[n] + yi_hours[n]

    def _weekday_eight_hours(n: str) -> int:
        return sum(
            8
            for dm in days
            if (not dm.is_weekend) and grid[n].get(dm.iso, "") == "8"
        )

    def _skip_weekend_night(n: str, iso: str, *, as_if_empty: bool = False) -> bool:
        """Hafta içi 8 ile mesaisi dolan / 8-pin → hafta sonu nöbet yok (pin 24 hariç).

        as_if_empty: mevcut hücre saati çıkarılır (strip için).
        """
        if (n, iso) in pinned_cells:
            return False
        # Özel koşul 8 pin: hafta sonu ek nöbet yok
        if n in relax_consec_8:
            return True
        # Yalnız hafta içi 8 (+Yİ) ideal'i dolduruyorsa ek mesai yazma
        if _weekday_eight_hours(n) + yi_hours[n] >= ideal:
            return True
        return False

    for idx, dm in enumerate(days):
        available = [
            n
            for n in STAFF_NURSES
            if not grid[n][dm.iso]
            and dm.iso not in force_avoid[n]
            and not _blocked_by_rest(n, idx, days, grid, prefer_48h_after_24=prefer_48h_after_24)
        ]
        leave_heavy = _week_is_leave_heavy(idx, days, grid)

        def is_gun_asiri_candidate(n: str) -> bool:
            return _gun_asiri_24_penalty(n, idx, days, grid) > 0

        def streak_if_24(n: str) -> int:
            return _gun_asiri_streak_if_24(n, idx, days, grid)

        def over_streak(n: str) -> bool:
            # Kati: 4. gün aşırı 24 yok (izin yoğun haftada da)
            return streak_if_24(n) > GUN_ASIRI_STREAK_MAX

        def special_work_rank(n: str) -> int:
            if dm.iso in prefer_work[n]:
                return 0
            if prefer_work[n]:
                return 2
            return 1

        def rank_for_8(n: str) -> tuple:
            pen = _rest_penalty(n, idx, days, grid)
            break_streak = 0 if over_streak(n) else 1
            break24 = 0 if is_gun_asiri_candidate(n) else 1
            # Özel koşul 8 pinli kişide art arda 8 cezası yok (birincil)
            if n in relax_consec_8:
                streak8 = 0
                prev8 = 0
            else:
                streak8 = _consecutive_8_streak_if_8(n, idx, days, grid)
                prev8 = 1 if streak8 >= 2 else 0
            after24 = _prefer_8_after_24_gap(n, idx, days, grid)
            # İzinden dönüş günü 8 değil nöbet (24)
            after_leave = 1 if _first_day_after_leave(n, idx, days, grid) else 0
            return (
                pen,
                after_leave,
                special_work_rank(n),
                streak8,
                break_streak,
                accounted(n),
                prev8,
                after24,
                n8[n],
                break24,
                _order(n),
                _tie(n),
                n,
            )

        def rank_for_24(n: str) -> tuple:
            pen = _rest_penalty(n, idx, days, grid)
            if n in day_only_set:
                pen += 500
            # Hafta sonu: mesaisi dolan / hafta içi-8 pin → nöbet son çare
            weekend_skip = 1 if (dm.is_weekend and _skip_weekend_night(n, dm.iso)) else 0
            over = 1 if over_streak(n) else 0
            gun = _gun_asiri_24_penalty(n, idx, days, grid)
            behind = 0 if hours[n] < min_shift[n] else 1
            # İzinden dönüş → önce nöbet
            after_leave = 0 if _first_day_after_leave(n, idx, days, grid) else 1
            return (
                weekend_skip,
                pen,
                after_leave,
                special_work_rank(n),
                over,
                accounted(n),
                streak_if_24(n),
                gun,
                behind,
                n24[n],
                _order(n),
                _tie(n),
                n,
            )

        morning: str | None = None
        night_needed = max(0, _night_target(dm.iso) - _staff_night_count(grid, dm.iso))

        def _assign24(n: str) -> None:
            nonlocal night_needed, available
            grid[n][dm.iso] = "24"
            hours[n] += 24
            n24[n] += 1
            night_needed -= 1
            available = [x for x in available if x != n]

        # Önce 2×24 nöbet (16 yok — sorumlu pinler)
        while night_needed > 0:
            pool = [
                n
                for n in available
                if n not in day_only_set
                and not _next_day_blocks_24(n, idx, days, grid)
                and not (dm.is_weekend and _skip_weekend_night(n, dm.iso))
            ]
            if not leave_heavy:
                capped = [n for n in pool if not over_streak(n)]
                if capped:
                    pool = capped
            if not pool:
                pool = [
                    n
                    for n in available
                    if n not in day_only_set
                    and not _next_day_blocks_24(n, idx, days, grid)
                    and not (dm.is_weekend and _skip_weekend_night(n, dm.iso))
                ]
            if not pool:
                break
            return_pool = [n for n in pool if _first_day_after_leave(n, idx, days, grid)]
            pick_from = return_pool if return_pool else pool
            _assign24(sorted(pick_from, key=rank_for_24)[0])

        # Hafta içi kat-1: her zaman bir «8» (nöbetlerden sonra; 8→24 çalma yok)
        if not dm.is_weekend:
            morning_cands = [
                n
                for n in sorted(available, key=rank_for_8)
                if n8[n] < EIGHT_PER_PERSON_MAX and _ok8(n, idx)
            ]
            no_return = [
                n for n in morning_cands if not _first_day_after_leave(n, idx, days, grid)
            ]
            if no_return:
                morning_cands = no_return
            if not morning_cands:
                # Max dolu olsa bile kat-1 için bir aday dene
                morning_cands = [
                    n
                    for n in sorted(available, key=rank_for_8)
                    if _ok8(n, idx)
                ]
                no_return = [
                    n
                    for n in morning_cands
                    if not _first_day_after_leave(n, idx, days, grid)
                ]
                if no_return:
                    morning_cands = no_return
            if morning_cands:
                morning = morning_cands[0]
                grid[morning][dm.iso] = "8"
                hours[morning] += 8
                n8[morning] += 1
                available = [n for n in available if n != morning]

        if dm.is_weekday:
            has_kat1 = any(grid[n][dm.iso] == "8" for n in STAFF_NURSES)
            if not has_kat1:
                warnings.append(f"{dm.iso}: Kat-1 gündüz 8 atanamadı (izin/dinlenme).")
        if night_needed > 0:
            tgt = _night_target(dm.iso)
            have = _staff_night_count(grid, dm.iso)
            warnings.append(f"{dm.iso}: Gece nöbeti eksik ({have}/{tgt}).")

    # ── Post: streak > 3 olan 24'leri mümkünse 8 ile takas (izin yoğun olsa da) ──
    for _pass in range(8):
        swapped = False
        for name in STAFF_NURSES:
            for i in range(len(days)):
                if grid[name].get(days[i].iso, "") != "24":
                    continue
                if (name, days[i].iso) in pinned_cells:
                    continue
                streak = _gun_asiri_streak_if_24(name, i, days, grid)
                soft_only = (
                    streak <= GUN_ASIRI_STREAK_MAX
                    and _gun_asiri_24_penalty(name, i, days, grid) > 0
                    and not _week_is_leave_heavy(i, days, grid)
                    and n8[name] < EIGHT_PER_PERSON_TARGET
                )
                hard = streak > GUN_ASIRI_STREAK_MAX
                if not hard and not soft_only:
                    continue
                if hard and n8[name] >= EIGHT_PER_PERSON_MAX and not _ok8(name, i):
                    # 8 yazılamazsa boşalt; gece doldurma sonra tamamlar
                    iso = days[i].iso
                    grid[name][iso] = ""
                    hours[name] -= 24
                    n24[name] -= 1
                    swapped = True
                    continue
                iso = days[i].iso
                partners = [
                    o
                    for o in STAFF_NURSES
                    if o != name
                    and grid[o].get(iso, "") == "8"
                    and (o, iso) not in pinned_cells
                    and not _cannot_assign_24(
                        o, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    )
                    and (
                        i + 1 >= len(days)
                        or grid[o].get(days[i + 1].iso, "") != "8"
                    )
                    and _gun_asiri_streak_if_24(o, i, days, grid) <= GUN_ASIRI_STREAK_MAX
                ]
                if not partners:
                    if hard and _ok8(name, i):
                        # Partner yok: 24→8 (gece eksiği sonraki pass doldurur)
                        grid[name][iso] = "8"
                        hours[name] -= 16
                        n24[name] -= 1
                        n8[name] += 1
                        swapped = True
                    elif hard:
                        grid[name][iso] = ""
                        hours[name] -= 24
                        n24[name] -= 1
                        swapped = True
                    continue
                other = sorted(partners, key=lambda o: (n24[o], hours[o], o))[0]
                if not _ok8(name, i):
                    if hard:
                        grid[name][iso] = ""
                        hours[name] -= 24
                        n24[name] -= 1
                        swapped = True
                    continue
                grid[name][iso] = "8"
                grid[other][iso] = "24"
                hours[name] -= 16
                hours[other] += 16
                n8[name] += 1
                n24[name] -= 1
                n8[other] -= 1
                n24[other] += 1
                swapped = True
        if not swapped:
            break

    # ── Fazla 8 budama (hedef üstü / max): gündüz 24 ile örtülüyse 8'i kaldır ──
    for name in STAFF_NURSES:
        while n8[name] > EIGHT_PER_PERSON_TARGET:
            trimmed = False
            for dm in days:
                if grid[name].get(dm.iso, "") != "8":
                    continue
                if (name, dm.iso) in pinned_cells:
                    continue
                day_24 = sum(1 for o in STAFF_NURSES if grid[o].get(dm.iso, "") == "24")
                if day_24 < 1:
                    continue
                grid[name][dm.iso] = ""
                hours[name] -= 8
                n8[name] -= 1
                trimmed = True
                break
            if not trimmed:
                break

    # Sert tavan 4 (örtü yoksa bile bırakmamak için son çare budama yok)
    for name in STAFF_NURSES:
        while n8[name] > EIGHT_PER_PERSON_MAX:
            trimmed = False
            for dm in days:
                if grid[name].get(dm.iso, "") != "8":
                    continue
                if (name, dm.iso) in pinned_cells:
                    continue
                day_24 = sum(1 for o in STAFF_NURSES if grid[o].get(dm.iso, "") == "24")
                if day_24 < 1:
                    continue
                grid[name][dm.iso] = ""
                hours[name] -= 8
                n8[name] -= 1
                trimmed = True
                break
            if not trimmed:
                break

    # ── Hafta sonu personel 8 temizle (kat-1 yok) — pin hariç ──
    for name in STAFF_NURSES:
        for dm in days:
            if not dm.is_weekend:
                continue
            if grid[name].get(dm.iso, "") != "8":
                continue
            if (name, dm.iso) in pinned_cells:
                continue
            grid[name][dm.iso] = ""
            hours[name] -= 8
            n8[name] -= 1

    # ── Eksik 8 (min 2): hafta içi; üst üste 8 yazma ──
    for name in STAFF_NURSES:
        for i, dm in enumerate(days):
            if n8[name] >= EIGHT_PER_PERSON_MIN:
                break
            if dm.is_weekend or grid[name].get(dm.iso, ""):
                continue
            if _blocked_by_rest(name, i, days, grid, prefer_48h_after_24=prefer_48h_after_24):
                continue
            if not _ok8(name, i):
                continue
            if n8[name] >= EIGHT_PER_PERSON_MAX:
                break
            grid[name][dm.iso] = "8"
            hours[name] += 8
            n8[name] += 1

    # ── Fazla mesai bandı: yüksek ↔ düşük (hedef ≤16s) ──
    def _strip_next_8_if_safe(recv: str, i: int) -> bool:
        """recv yarın 8 ise ve gündüz başka örtü varsa 8'i kaldır → 24 alabilsin."""
        if i + 1 >= len(days):
            return True
        nxt = days[i + 1]
        code = grid[recv].get(nxt.iso, "")
        if code in ("16", "24"):
            return False
        if code != "8":
            return True
        if (recv, nxt.iso) in pinned_cells:
            return False
        other_morn = any(
            o != recv and grid[o].get(nxt.iso, "") in ("8", "24") for o in STAFF_NURSES
        )
        if not other_morn and not nxt.is_weekend:
            return False
        grid[recv][nxt.iso] = ""
        hours[recv] -= 8
        n8[recv] -= 1
        return True

    def _can_take_24(recv: str, i: int) -> bool:
        if grid[recv].get(days[i].iso, ""):
            return False
        if recv in day_only_set:
            return False
        if days[i].is_weekend and _skip_weekend_night(recv, days[i].iso):
            return False
        if _cannot_assign_24(
            recv, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
        ):
            return False
        if (not _week_is_leave_heavy(i, days, grid)) and (
            _gun_asiri_streak_if_24(recv, i, days, grid) > GUN_ASIRI_STREAK_MAX
        ):
            return False
        if not _strip_next_8_if_safe(recv, i):
            return False
        return True

    for _bal in range(80):
        vals = {n: accounted(n) for n in STAFF_NURSES}
        hi = max(STAFF_NURSES, key=lambda n: (vals[n], _tie(n), n))
        lo = min(STAFF_NURSES, key=lambda n: (vals[n], _tie(n), n))
        if vals[hi] - vals[lo] <= HOURS_BALANCE_TOLERANCE:
            break
        moved = False

        # Aynı gün hi=24 lo=8 → takas
        for i in day_order:
            dm = days[i]
            if grid[hi].get(dm.iso, "") != "24":
                continue
            if (hi, dm.iso) in pinned_cells:
                continue
            if grid[lo].get(dm.iso, "") != "8":
                continue
            if (lo, dm.iso) in pinned_cells:
                continue
            if _cannot_assign_24(
                lo, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
            ):
                continue
            if n8[hi] >= EIGHT_PER_PERSON_MAX:
                continue
            if not _ok8(hi, i):
                continue
            # Takas sonrası lo'nun streak'i (bugün 24 olacak)
            if (not _week_is_leave_heavy(i, days, grid)) and (
                _gun_asiri_streak_if_24(lo, i, days, grid) > GUN_ASIRI_STREAK_MAX
            ):
                continue
            if not _strip_next_8_if_safe(lo, i):
                continue
            grid[hi][dm.iso] = "8"
            grid[lo][dm.iso] = "24"
            hours[hi] -= 16
            hours[lo] += 16
            n24[hi] -= 1
            n24[lo] += 1
            n8[hi] += 1
            n8[lo] -= 1
            moved = True
            break
        if moved:
            continue

        # hi 24 → lo boş
        for i in day_order:
            dm = days[i]
            if grid[hi].get(dm.iso, "") != "24":
                continue
            if (hi, dm.iso) in pinned_cells:
                continue
            if not _can_take_24(lo, i):
                continue
            grid[hi][dm.iso] = ""
            grid[lo][dm.iso] = "24"
            hours[hi] -= 24
            hours[lo] += 24
            n24[hi] -= 1
            n24[lo] += 1
            moved = True
            break
        if moved:
            continue

        # hi 8 → lo boş (hafta içi)
        for i in day_order:
            dm = days[i]
            if dm.is_weekend:
                continue
            if grid[hi].get(dm.iso, "") != "8":
                continue
            if (hi, dm.iso) in pinned_cells:
                continue
            if grid[lo].get(dm.iso, ""):
                continue
            if _blocked_by_rest(lo, i, days, grid, prefer_48h_after_24=prefer_48h_after_24):
                continue
            if n8[lo] >= EIGHT_PER_PERSON_MAX or n8[hi] <= EIGHT_PER_PERSON_MIN:
                continue
            if not _ok8(lo, i):
                continue
            grid[hi][dm.iso] = ""
            grid[lo][dm.iso] = "8"
            hours[hi] -= 8
            hours[lo] += 8
            n8[hi] -= 1
            n8[lo] += 1
            moved = True
            break
        if moved:
            continue

        # hi=24 mid=8
        mids = [
            n
            for n in sorted(STAFF_NURSES, key=lambda n: (vals[n], _tie(n), n))
            if n not in (hi, lo)
        ]
        for mid in mids:
            for i in day_order:
                dm = days[i]
                if grid[hi].get(dm.iso, "") != "24":
                    continue
                if (hi, dm.iso) in pinned_cells:
                    continue
                if grid[mid].get(dm.iso, "") != "8":
                    continue
                if (mid, dm.iso) in pinned_cells:
                    continue
                if _cannot_assign_24(
                    mid, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                ):
                    continue
                if n8[hi] >= EIGHT_PER_PERSON_MAX:
                    continue
                if not _ok8(hi, i):
                    continue
                if (not _week_is_leave_heavy(i, days, grid)) and (
                    _gun_asiri_streak_if_24(mid, i, days, grid) > GUN_ASIRI_STREAK_MAX
                ):
                    continue
                if not _strip_next_8_if_safe(mid, i):
                    continue
                grid[hi][dm.iso] = "8"
                grid[mid][dm.iso] = "24"
                hours[hi] -= 16
                hours[mid] += 16
                n24[hi] -= 1
                n24[mid] += 1
                n8[hi] += 1
                n8[mid] -= 1
                moved = True
                break
            if moved:
                break
        if not moved:
            break

    # Denge sonrası: min 8 ve hafta sonu 8 temizliği
    for name in STAFF_NURSES:
        for dm in days:
            if dm.is_weekend and grid[name].get(dm.iso, "") == "8":
                grid[name][dm.iso] = ""
                hours[name] -= 8
                n8[name] -= 1
    for name in STAFF_NURSES:
        for i, dm in enumerate(days):
            if n8[name] >= EIGHT_PER_PERSON_MIN:
                break
            if dm.is_weekend or grid[name].get(dm.iso, ""):
                continue
            if _blocked_by_rest(name, i, days, grid, prefer_48h_after_24=prefer_48h_after_24):
                continue
            if not _ok8(name, i):
                continue
            grid[name][dm.iso] = "8"
            hours[name] += 8
            n8[name] += 1

    # Streak > 3 kır: zincirdeki son 24'ü başka kişiye ver veya 8 yap
    for _sk in range(20):
        broke = False
        for name in STAFF_NURSES:
            for i in range(len(days)):
                if grid[name].get(days[i].iso, "") != "24":
                    continue
                if _week_is_leave_heavy(i, days, grid):
                    continue
                if _gun_asiri_streak_if_24(name, i, days, grid) <= GUN_ASIRI_STREAK_MAX:
                    continue
                # Bu gün zinciri uzatıyor — taşı veya 8'e düş
                iso = days[i].iso
                if (name, iso) in pinned_cells:
                    continue
                takers = [
                    o
                    for o in STAFF_NURSES
                    if o != name
                    and grid[o].get(iso, "") == ""
                    and o not in day_only_set
                    and not _cannot_assign_24(
                        o, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    )
                    and _gun_asiri_streak_if_24(o, i, days, grid) <= GUN_ASIRI_STREAK_MAX
                ]
                if takers:
                    other = sorted(takers, key=lambda o: (accounted(o), n24[o], o))[0]
                    grid[name][iso] = ""
                    grid[other][iso] = "24"
                    hours[name] -= 24
                    hours[other] += 24
                    n24[name] -= 1
                    n24[other] += 1
                    broke = True
                    break
                # 24→8 düşürme: günde zaten 2. gece varken kapsamayı 1'e indirir — yapma
            if broke:
                break
        if not broke:
            break

    # Streak kırımı sonrası tekrar kısa denge
    for _bal2 in range(40):
        vals = {n: accounted(n) for n in STAFF_NURSES}
        if max(vals.values()) - min(vals.values()) <= HOURS_BALANCE_TOLERANCE:
            break
        moved = False
        his = sorted(STAFF_NURSES, key=lambda n: (-vals[n], -n24[n], n))
        los = sorted(STAFF_NURSES, key=lambda n: (vals[n], n24[n], n))
        for hi in his:
            for lo in los:
                if hi == lo or vals[hi] - vals[lo] <= HOURS_BALANCE_TOLERANCE:
                    continue
                for i, dm in enumerate(days):
                    if grid[hi].get(dm.iso, "") != "24":
                        continue
                    if (hi, dm.iso) in pinned_cells:
                        continue
                    if not _can_take_24(lo, i):
                        continue
                    grid[hi][dm.iso] = ""
                    grid[lo][dm.iso] = "24"
                    hours[hi] -= 24
                    hours[lo] += 24
                    n24[hi] -= 1
                    n24[lo] += 1
                    moved = True
                    break
                if moved:
                    break
                for i, dm in enumerate(days):
                    if grid[hi].get(dm.iso, "") != "24" or grid[lo].get(dm.iso, "") != "8":
                        continue
                    if (hi, dm.iso) in pinned_cells:
                        continue
                    if (lo, dm.iso) in pinned_cells:
                        continue
                    if _cannot_assign_24(
                        lo, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    ):
                        continue
                    if n8[hi] >= EIGHT_PER_PERSON_MAX:
                        continue
                    if (not _week_is_leave_heavy(i, days, grid)) and (
                        _gun_asiri_streak_if_24(lo, i, days, grid) > GUN_ASIRI_STREAK_MAX
                    ):
                        continue
                    if not _strip_next_8_if_safe(lo, i):
                        continue
                    if not _ok8(hi, i):
                        continue
                    grid[hi][dm.iso] = "8"
                    grid[lo][dm.iso] = "24"
                    hours[hi] -= 16
                    hours[lo] += 16
                    n24[hi] -= 1
                    n24[lo] += 1
                    n8[hi] += 1
                    n8[lo] -= 1
                    moved = True
                    break
                if moved:
                    break
            if moved:
                break
        if not moved:
            break

    # Üst üste 3+ «8» kır (mecburen en fazla 2) — özel koşul 8 pinli kişi muaf
    for _fix8long in range(20):
        fixed = False
        for name in STAFF_NURSES:
            if name in relax_consec_8:
                continue
            if _max_consecutive_8_streak(name, days, grid) <= CONSECUTIVE_8_STREAK_MAX:
                continue
            target_i: int | None = None
            run_start: int | None = None
            for i, dm in enumerate(days):
                if grid[name].get(dm.iso, "") == "8":
                    if run_start is None:
                        run_start = i
                    if i - run_start + 1 > CONSECUTIVE_8_STREAK_MAX:
                        target_i = i
                else:
                    run_start = None
            if target_i is None:
                continue
            iso = days[target_i].iso
            if (name, iso) in pinned_cells:
                continue
            covered = any(
                o != name and grid[o].get(iso, "") in ("8", "24") for o in STAFF_NURSES
            )
            if covered:
                grid[name][iso] = ""
                hours[name] -= 8
                n8[name] -= 1
                fixed = True
                continue
            moved = False
            for j, dm in enumerate(days):
                if dm.is_weekend or grid[name].get(dm.iso, ""):
                    continue
                if (name, dm.iso) in pinned_cells:
                    continue
                if not _ok8(name, j):
                    continue
                if _blocked_by_rest(
                    name, j, days, grid, prefer_48h_after_24=prefer_48h_after_24
                ):
                    continue
                if _prefer_8_after_24_gap(name, j, days, grid) > 0:
                    continue
                grid[name][iso] = ""
                grid[name][dm.iso] = "8"
                fixed = True
                moved = True
                break
            if not moved and not covered:
                leave_heavy_day = _week_is_leave_heavy(target_i, days, grid)
                if (
                    name not in day_only_set
                    and (name, iso) not in pinned_cells
                    and not _cannot_assign_24(
                        name, target_i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    )
                    and (
                        leave_heavy_day
                        or _gun_asiri_streak_if_24(name, target_i, days, grid)
                        <= GUN_ASIRI_STREAK_MAX
                    )
                ):
                    grid[name][iso] = "24"
                    hours[name] += 16
                    n8[name] -= 1
                    n24[name] += 1
                    fixed = True
        if not fixed:
            break

    # Son: kalan üst üste 8 → yalnız 24 arasına kaydır (saat silme)
    for _e8f in range(8):
        fixed = False
        for name in STAFF_NURSES:
            if name in relax_consec_8:
                continue
            for i in range(1, len(days)):
                iso = days[i].iso
                if grid[name].get(iso, "") != "8":
                    continue
                if (name, iso) in pinned_cells:
                    continue
                if grid[name].get(days[i - 1].iso, "") != "8":
                    continue
                for j in range(2, len(days)):
                    if days[j].is_weekend or grid[name].get(days[j].iso, ""):
                        continue
                    if _blocked_by_rest(
                        name, j, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    ):
                        continue
                    if _prev_day_is_8(name, j, days, grid):
                        continue
                    if not _ok8(name, j):
                        continue
                    if _prefer_8_after_24_gap(name, j, days, grid) > 0:
                        continue
                    if grid[name].get(days[j - 2].iso, "") != "24":
                        continue
                    if grid[name].get(days[j - 1].iso, "") not in ("", "Yİ", "RP", "İST"):
                        continue
                    grid[name][iso] = ""
                    grid[name][days[j].iso] = "8"
                    fixed = True
                    break
                if fixed:
                    break
            if fixed:
                break
        if not fixed:
            break

    def _can_steal_8_for_24(n: str, day_index: int, iso: str) -> bool:
        """Kat-1 tek 8'i çalma; pin'e dokunma; hafta sonu veya başka 8 varken izin ver."""
        if grid[n].get(iso, "") != "8":
            return False
        if (n, iso) in pinned_cells:
            return False
        if days[day_index].is_weekend:
            return True
        return any(
            o != n and grid[o].get(iso, "") == "8" for o in STAFF_NURSES
        )

    # ── Zorunlu: her gün tam 2× gece nöbeti (post-pass sonrası boşluk kalmasın) ──
    for idx, dm in enumerate(days):
        iso = dm.iso
        while _staff_night_count(grid, iso) < _night_target(iso):
            filled = False

            def _rank_night_fill(n: str) -> tuple:
                code = grid[n].get(iso, "")
                empty = 0 if code == "" else 1
                over = max(
                    0,
                    _gun_asiri_streak_if_24(n, idx, days, grid) - GUN_ASIRI_STREAK_MAX,
                )
                return (empty, over, accounted(n), n24[n], _order(n), _tie(n), n)

            # Kati: streak ≤3; aşacak aday yoksa streak gevşet — 16 yok
            pool = [
                n
                for n in STAFF_NURSES
                if n not in day_only_set
                and grid[n].get(iso, "") not in LEAVE_CODES
                and grid[n].get(iso, "") == ""
                and not (dm.is_weekend and _skip_weekend_night(n, iso))
                and not _cannot_assign_24(
                    n, idx, days, grid, prefer_48h_after_24=prefer_48h_after_24
                )
                and _gun_asiri_streak_if_24(n, idx, days, grid) <= GUN_ASIRI_STREAK_MAX
            ]
            if not pool:
                pool = [
                    n
                    for n in STAFF_NURSES
                    if n not in day_only_set
                    and grid[n].get(iso, "") not in LEAVE_CODES
                    and grid[n].get(iso, "") == ""
                    and not (dm.is_weekend and _skip_weekend_night(n, iso))
                    and not _cannot_assign_24(
                        n, idx, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    )
                ]
            # Son çare: hafta sonu veya kat-1 başka 8 varsa 8→24
            if not pool:
                pool = [
                    n
                    for n in STAFF_NURSES
                    if n not in day_only_set
                    and _can_steal_8_for_24(n, idx, iso)
                    and not (dm.is_weekend and _skip_weekend_night(n, iso))
                    and not _cannot_assign_24(
                        n, idx, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    )
                ]
            if pool:
                pick = sorted(pool, key=_rank_night_fill)[0]
                prev = grid[pick].get(iso, "")
                grid[pick][iso] = "24"
                if prev == "8":
                    hours[pick] += 16
                    n8[pick] -= 1
                else:
                    hours[pick] += 24
                n24[pick] += 1
                filled = True

            if filled:
                continue

            warnings.append(
                f"{iso}: Gece nöbeti eksik ({_staff_night_count(grid, iso)}/{_night_target(iso)})."
            )
            break

    # Son kilit: hiçbir personelde 4+ gün aşırı 24 kalmasın
    for _ in range(16):
        offender: tuple[str, int] | None = None
        worst = 0
        for name in STAFF_NURSES:
            for i in range(len(days)):
                if grid[name].get(days[i].iso, "") != "24":
                    continue
                if (name, days[i].iso) in pinned_cells:
                    continue
                streak = _gun_asiri_streak_if_24(name, i, days, grid)
                if streak > GUN_ASIRI_STREAK_MAX and streak > worst:
                    worst = streak
                    offender = (name, i)
        if not offender:
            break
        name, i = offender
        iso = days[i].iso
        grid[name][iso] = ""
        hours[name] -= 24
        n24[name] -= 1
        # Gün gece eksiği: yalnız 24 (16 yok)
        while _staff_night_count(grid, iso) < _night_target(iso):
            pool = [
                n
                for n in STAFF_NURSES
                if n != name
                and n not in day_only_set
                and grid[n].get(iso, "") == ""
                and not (days[i].is_weekend and _skip_weekend_night(n, iso))
                and not _cannot_assign_24(
                    n, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                )
                and _gun_asiri_streak_if_24(n, i, days, grid) <= GUN_ASIRI_STREAK_MAX
            ]
            if not pool:
                pool = [
                    n
                    for n in STAFF_NURSES
                    if n != name
                    and n not in day_only_set
                    and grid[n].get(iso, "") == ""
                    and not (days[i].is_weekend and _skip_weekend_night(n, iso))
                    and not _cannot_assign_24(
                        n, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    )
                ]
            if not pool:
                pool = [
                    n
                    for n in STAFF_NURSES
                    if n != name
                    and n not in day_only_set
                    and _can_steal_8_for_24(n, i, iso)
                    and not (days[i].is_weekend and _skip_weekend_night(n, iso))
                    and not _cannot_assign_24(
                        n, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    )
                ]
            if pool:
                pick = sorted(
                    pool, key=lambda n: (accounted(n), n24[n], _order(n), _tie(n), n)
                )[0]
                prev = grid[pick].get(iso, "")
                grid[pick][iso] = "24"
                if prev == "8":
                    hours[pick] += 16
                    n8[pick] -= 1
                else:
                    hours[pick] += 24
                n24[pick] += 1
                continue
            break

    # Son kilit: aynı kişide arka arkaya 24 yasak (pin hariç)
    for _adj in range(24):
        conflict: tuple[str, int] | None = None
        for name in STAFF_NURSES:
            for i in range(len(days) - 1):
                if grid[name].get(days[i].iso, "") != "24":
                    continue
                if grid[name].get(days[i + 1].iso, "") != "24":
                    continue
                # Tercihen ikinci günü taşı (unpin)
                if (name, days[i + 1].iso) not in pinned_cells:
                    conflict = (name, i + 1)
                elif (name, days[i].iso) not in pinned_cells:
                    conflict = (name, i)
                else:
                    continue
                break
            if conflict:
                break
        if not conflict:
            break
        name, i = conflict
        iso = days[i].iso
        grid[name][iso] = ""
        hours[name] -= 24
        n24[name] -= 1
        while _staff_night_count(grid, iso) < _night_target(iso):
            pool = [
                n
                for n in STAFF_NURSES
                if n != name
                and n not in day_only_set
                and grid[n].get(iso, "") == ""
                and not (days[i].is_weekend and _skip_weekend_night(n, iso))
                and not _cannot_assign_24(
                    n, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                )
                and _gun_asiri_streak_if_24(n, i, days, grid) <= GUN_ASIRI_STREAK_MAX
            ]
            if not pool:
                pool = [
                    n
                    for n in STAFF_NURSES
                    if n != name
                    and n not in day_only_set
                    and grid[n].get(iso, "") == ""
                    and not (days[i].is_weekend and _skip_weekend_night(n, iso))
                    and not _cannot_assign_24(
                        n, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    )
                ]
            if not pool:
                pool = [
                    n
                    for n in STAFF_NURSES
                    if n != name
                    and n not in day_only_set
                    and _can_steal_8_for_24(n, i, iso)
                    and not (days[i].is_weekend and _skip_weekend_night(n, iso))
                    and not _cannot_assign_24(
                        n, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    )
                ]
            if pool:
                pick = sorted(
                    pool, key=lambda n: (accounted(n), n24[n], _order(n), _tie(n), n)
                )[0]
                prev = grid[pick].get(iso, "")
                grid[pick][iso] = "24"
                if prev == "8":
                    hours[pick] += 16
                    n8[pick] -= 1
                else:
                    hours[pick] += 24
                n24[pick] += 1
                continue
            break

    # Hafta sonu: mesaisi dolan / hafta içi-8 pin kişiden nöbet kaldır (ek mesai yok)
    for idx, dm in enumerate(days):
        if not dm.is_weekend:
            continue
        for name in STAFF_NURSES:
            code = grid[name].get(dm.iso, "")
            if code not in ("16", "24"):
                continue
            if (name, dm.iso) in pinned_cells:
                continue
            if not _skip_weekend_night(name, dm.iso, as_if_empty=True):
                continue
            hours[name] -= _hours_for(code)
            if code == "24":
                n24[name] -= 1
            else:
                n16[name] -= 1
            grid[name][dm.iso] = ""

    # Motor 16 yazmaz: pin olmayan 16 → mümkünse 24, değilse boş
    for name in STAFF_NURSES:
        for i, dm in enumerate(days):
            if grid[name].get(dm.iso, "") != "16":
                continue
            if (name, dm.iso) in pinned_cells:
                continue
            grid[name][dm.iso] = ""
            hours[name] -= 16
            n16[name] -= 1
            if dm.is_weekend and _skip_weekend_night(name, dm.iso):
                continue
            if (
                name not in day_only_set
                and not _cannot_assign_24(
                    name, i, days, grid, prefer_48h_after_24=prefer_48h_after_24
                )
            ):
                grid[name][dm.iso] = "24"
                hours[name] += 24
                n24[name] += 1

    # Hafta içi kat-1 8 son geçiş (post-pass sonrası eksik kalmasın)
    for idx, dm in enumerate(days):
        if dm.is_weekend:
            continue
        if any(grid[n].get(dm.iso, "") == "8" for n in STAFF_NURSES):
            continue
        cands = [
            n
            for n in sorted(STAFF_NURSES, key=lambda n: (n8[n], accounted(n), _tie(n), n))
            if not grid[n].get(dm.iso, "")
            and dm.iso not in force_avoid[n]
            and not _blocked_by_rest(
                n, idx, days, grid, prefer_48h_after_24=prefer_48h_after_24
            )
            and _ok8(n, idx)
            and not _first_day_after_leave(n, idx, days, grid)
        ]
        if not cands:
            cands = [
                n
                for n in sorted(STAFF_NURSES, key=lambda n: (n8[n], accounted(n), _tie(n), n))
                if not grid[n].get(dm.iso, "")
                and dm.iso not in force_avoid[n]
                and not _blocked_by_rest(
                    n, idx, days, grid, prefer_48h_after_24=prefer_48h_after_24
                )
                and _ok8(n, idx)
            ]
        if cands:
            pick = cands[0]
            grid[pick][dm.iso] = "8"
            hours[pick] += 8
            n8[pick] += 1

    # Son gece kilidi: her gün 2×24 (16 yok); pin koru; gerekirse yarınki 24'ü kaydır
    for idx, dm in enumerate(days):
        iso = dm.iso
        guard = 0
        while _staff_night_count(grid, iso) < _night_target(iso) and guard < 8:
            guard += 1

            def _try_place_24(n: str) -> bool:
                if n in day_only_set or (n, iso) in pinned_cells:
                    return False
                if dm.is_weekend and _skip_weekend_night(n, iso):
                    return False
                code = grid[n].get(iso, "")
                if code in LEAVE_CODES or code in ("16", "24"):
                    return False
                if code == "8" and (n, iso) in pinned_cells:
                    return False
                if _blocked_by_rest(
                    n, idx, days, grid, prefer_48h_after_24=prefer_48h_after_24
                ):
                    return False
                # Yarın 24 ise (pin değilse) kaydırıp bugün yaz
                if _next_day_blocks_24(n, idx, days, grid):
                    nxt = days[idx + 1]
                    if (n, nxt.iso) in pinned_cells:
                        return False
                    if grid[n].get(nxt.iso, "") != "24":
                        return False
                    # Yarını boşalt; gece eksiği sonraki turda dolar
                    grid[n][nxt.iso] = ""
                    hours[n] -= 24
                    n24[n] -= 1
                if code == "8":
                    grid[n][iso] = "24"
                    hours[n] += 16
                    n8[n] -= 1
                    n24[n] += 1
                    return True
                if code == "":
                    grid[n][iso] = "24"
                    hours[n] += 24
                    n24[n] += 1
                    return True
                return False

            placed = False
            # 1) boş + streak güvenli
            pool = sorted(
                [
                    n
                    for n in STAFF_NURSES
                    if grid[n].get(iso, "") == ""
                    and not _cannot_assign_24(
                        n, idx, days, grid, prefer_48h_after_24=prefer_48h_after_24
                    )
                    and _gun_asiri_streak_if_24(n, idx, days, grid) <= GUN_ASIRI_STREAK_MAX
                ],
                key=lambda n: (accounted(n), n24[n], _tie(n), n),
            )
            for n in pool:
                if _try_place_24(n):
                    placed = True
                    break
            if placed:
                continue
            # 2) boş (streak gevşek)
            pool = sorted(
                [n for n in STAFF_NURSES if grid[n].get(iso, "") == ""],
                key=lambda n: (accounted(n), n24[n], _tie(n), n),
            )
            for n in pool:
                if _try_place_24(n):
                    placed = True
                    break
            if placed:
                continue
            # 3) 8 çal (pin hariç) — kat-1 sonra yeniden yazılır
            pool = sorted(
                [
                    n
                    for n in STAFF_NURSES
                    if grid[n].get(iso, "") == "8" and (n, iso) not in pinned_cells
                ],
                key=lambda n: (accounted(n), n24[n], _tie(n), n),
            )
            for n in pool:
                if _try_place_24(n):
                    placed = True
                    break
            if placed:
                continue
            warnings.append(
                f"{iso}: Gece nöbeti eksik ({_staff_night_count(grid, iso)}/{_night_target(iso)})."
            )
            break

    # Kat-1 8'i gece kilidinden sonra bir kez daha
    for idx, dm in enumerate(days):
        if dm.is_weekend:
            continue
        if any(grid[n].get(dm.iso, "") == "8" for n in STAFF_NURSES):
            continue
        cands = [
            n
            for n in sorted(STAFF_NURSES, key=lambda n: (n8[n], accounted(n), _tie(n), n))
            if not grid[n].get(dm.iso, "")
            and dm.iso not in force_avoid[n]
            and not _blocked_by_rest(
                n, idx, days, grid, prefer_48h_after_24=prefer_48h_after_24
            )
            and _ok8(n, idx)
        ]
        if cands:
            pick = cands[0]
            grid[pick][dm.iso] = "8"
            hours[pick] += 8
            n8[pick] += 1

    # Sabit pin son kilit — hiçbir post-pass ezemesin
    _apply_special_pins(grid, pins, hours, n8, n16, n24)
    yi_hours = {n: _yi_hours_from_grid(grid, n, days) for n in STAFF_NURSES}

    last = days[-1]
    next_month_rest = [
        n for n in STAFF_NURSES if grid[n].get(last.iso, "") in ("16", "24")
    ]

    accounted_list = [hours[n] + yi_hours[n] for n in STAFF_NURSES]
    if accounted_list:
        spread = max(accounted_list) - min(accounted_list)
        if spread > HOURS_BALANCE_TOLERANCE:
            warnings.append(
                f"Personel fazla/toplam saat bandı {spread:.0f}s "
                f"(hedef ≤{HOURS_BALANCE_TOLERANCE}s aralık)."
            )

    gun_asiri = 0
    for name in STAFF_NURSES:
        for i in range(2, len(days)):
            if grid[name].get(days[i].iso, "") != "24":
                continue
            if grid[name].get(days[i - 2].iso, "") != "24":
                continue
            mid = grid[name].get(days[i - 1].iso, "")
            if mid in ("", "Yİ", "RP", "İST"):
                gun_asiri += 1
    if gun_asiri:
        warnings.append(
            f"Gün aşırı 24 kalıbı {gun_asiri} kez (zincir tavanı ≤{GUN_ASIRI_STREAK_MAX}; "
            "4+ yasak)."
        )

    back_to_back = 0
    for name in STAFF_NURSES:
        for i in range(len(days) - 1):
            if (
                grid[name].get(days[i].iso, "") == "24"
                and grid[name].get(days[i + 1].iso, "") == "24"
            ):
                back_to_back += 1
    if back_to_back:
        warnings.append(
            f"Arka arkaya 24 kaldı: {back_to_back} çift (çoğunlukla sabit pin)."
        )

    rows: list[dict[str, Any]] = []
    for name in ALL_NURSES:
        is_lead = name == LEAD_NURSE
        if is_lead:
            shift_h = sum(_hours_for(grid[name][dm.iso]) for dm in days)
            leave_h = 0
            acc = shift_h
            target = 0
            min_s = 0
            overtime = 0
        else:
            shift_h = hours[name]
            leave_h = yi_hours[name]
            acc = shift_h + leave_h  # Yİ günleri 8s gibi yazılır
            target = ideal
            min_s = min_shift[name]
            overtime = acc - target  # eksi = eksik, artı = fazla
            if shift_h < min_s:
                warnings.append(
                    f"{name}: zorunlu mesai eksiği — en az {min_s}s nöbet "
                    f"(kota {ideal} − Yİ {leave_h}), şu an {shift_h}s."
                )
        rows.append(
            {
                "name": name,
                "role": "lead" if is_lead else "staff",
                "day_only": name in day_only_set,
                "cells": {dm.iso: grid[name][dm.iso] for dm in days},
                "worked_hours": acc,
                "shift_hours": shift_h,
                "leave_hours": leave_h,
                "ideal_hours": target,
                "min_shift_hours": min_s,
                "overtime_hours": overtime,
                "over_cap": (not is_lead) and shift_h > MAX_MONTHLY_HOURS,
                "exclude_from_staff_balance": is_lead,
                "count_8": 0 if is_lead else n8[name],
                "count_24": 0 if is_lead else n24[name],
                "count_16": 0 if is_lead else n16[name],
            }
        )

    jin_cells = {
        dm.iso: _format_jin_cell(jin_coverage.get(dm.iso, [])) for dm in days
    }
    if any(jin_cells.values()):
        jin_hours = sum(
            _hours_for(code)
            for slots in jin_coverage.values()
            for code in slots
        )
        rows.append(
            {
                "name": JIN_NURSE_LABEL,
                "role": "jin",
                "day_only": False,
                "cells": jin_cells,
                "worked_hours": jin_hours,
                "shift_hours": jin_hours,
                "leave_hours": 0,
                "ideal_hours": 0,
                "min_shift_hours": 0,
                "overtime_hours": 0,
                "over_cap": False,
                "exclude_from_staff_balance": True,
                "count_8": 0,
                "count_24": sum(
                    1
                    for slots in jin_coverage.values()
                    for code in slots
                    if code == "24"
                ),
                "count_16": sum(
                    1
                    for slots in jin_coverage.values()
                    for code in slots
                    if code == "16"
                ),
            }
        )

    code_counts = {"8": sum(n8.values()), "16": sum(n16.values()), "24": sum(n24.values())}

    return {
        "ok": True,
        "year": year,
        "month": month,
        "variant": variant_i,
        "month_label": f"{calendar.month_name[month]} {year}",
        "days": [
            {
                "day": dm.day,
                "iso": dm.iso,
                "weekday": dm.weekday,
                "weekday_tr": ("Pzt", "Sal", "Çar", "Per", "Cum", "Cmt", "Paz")[dm.weekday],
                "is_weekend": dm.is_weekend,
            }
            for dm in days
        ],
        "nurses": list(ALL_NURSES),
        "staff": list(STAFF_NURSES),
        "lead": LEAD_NURSE,
        "ideal_hours_staff": ideal,
        "weekday_count": count_weekdays(year, month),
        "max_monthly_hours": MAX_MONTHLY_HOURS,
        "hours_balance_tolerance": HOURS_BALANCE_TOLERANCE,
        "yi_day_hours": YI_DAY_HOURS,
        "rows": rows,
        "warnings": warnings,
        "next_month_must_rest": next_month_rest,
        "staff_code_counts": code_counts,
        "jin_coverage": {
            iso: list(slots) for iso, slots in sorted(jin_coverage.items())
        },
        "legend": {
            "8": "08:00–16:00 (6 kişiye dağıtılır)",
            "16": "16:00–08:00 (yalnız sabit pin)",
            "24": "08:00–08:00 (nöbet — motor her zaman 24)",
            "Yİ": f"Yıllık izin ({YI_DAY_HOURS}s sayılır)",
            "RP": "Rapor",
            "İST": "Özel gün isteği / rezervasyon",
        },
    }


def roster_defaults() -> dict[str, Any]:
    today = date.today()
    if today.month == 12:
        default_year, default_month = today.year + 1, 1
    else:
        default_year, default_month = today.year, today.month + 1
    return {
        "lead": LEAD_NURSE,
        "staff": list(STAFF_NURSES),
        "all": list(ALL_NURSES),
        "default_year": default_year,
        "default_month": default_month,
        "max_monthly_hours": MAX_MONTHLY_HOURS,
        "legend": {
            "8": "08:00–16:00",
            "16": "16:00–08:00",
            "24": "08:00–08:00",
            "Yİ": "Yıllık izin",
            "RP": "Rapor",
            "İST": "Özel gün isteği",
        },
    }


def build_ayilma_xlsx_bytes(
    *,
    year: int,
    month: int,
    days: list[dict[str, Any]],
    rows: list[dict[str, Any]],
) -> bytes:
    """Win + Mac Excel / Numbers / LibreOffice uyumlu .xlsx."""
    from io import BytesIO

    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    title = _month_title(month, year)
    export_rows = _ensure_empty_lead_export_rows(days, rows)
    headers, body = _export_matrix(days, export_rows)

    wb = Workbook()
    ws = wb.active
    ws.title = f"{month:02d}-{year}"[:31]

    ws["A1"] = title
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(4, 1 + len(days) + 3))

    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=3, column=col, value=h)
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    # Hafta sonu başlık boyası
    wknd_fill = PatternFill("solid", fgColor="BBF7D0")
    for i, d in enumerate(days):
        if d.get("is_weekend"):
            ws.cell(row=3, column=2 + i).fill = wknd_fill

    fills = {
        "Yİ": PatternFill("solid", fgColor="FDE68A"),
        "RP": PatternFill("solid", fgColor="FECACA"),
        "İST": PatternFill("solid", fgColor="BAE6FD"),
        "8": PatternFill("solid", fgColor="E0E7FF"),
        "16": PatternFill("solid", fgColor="FEF3C7"),
        "24": PatternFill("solid", fgColor="DDD6FE"),
    }

    for r_i, line in enumerate(body, start=4):
        row_meta = export_rows[r_i - 4] if r_i - 4 < len(export_rows) else {}
        ws.cell(row=r_i, column=1, value=line[0]).font = Font(bold=True)
        for c_i, d in enumerate(days):
            code = str(line[1 + c_i] or "")
            cell = ws.cell(row=r_i, column=2 + c_i, value=code)
            cell.alignment = Alignment(horizontal="center")
            if code in fills:
                cell.fill = fills[code]
        ws.cell(row=r_i, column=2 + len(days), value=line[-3])
        ws.cell(row=r_i, column=3 + len(days), value=line[-2])
        ws.cell(row=r_i, column=4 + len(days), value=line[-1])
        if row_meta.get("role") == "lead":
            ws.cell(row=r_i, column=1).font = Font(bold=True, color="4338CA")

    ws.column_dimensions["A"].width = 18
    for i in range(len(days)):
        ws.column_dimensions[ws.cell(row=3, column=2 + i).column_letter].width = 4.2
    for j in range(3):
        ws.column_dimensions[ws.cell(row=3, column=2 + len(days) + j).column_letter].width = 10

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


_MONTH_TR_NAMES = (
    "",
    "Ocak",
    "Şubat",
    "Mart",
    "Nisan",
    "Mayıs",
    "Haziran",
    "Temmuz",
    "Ağustos",
    "Eylül",
    "Ekim",
    "Kasım",
    "Aralık",
)


def _month_title(month: int, year: int) -> str:
    label = _MONTH_TR_NAMES[month] if 1 <= month <= 12 else str(month)
    return f"Ayılma hemşireleri — {label} {year}"


def _export_headers(days: list[dict[str, Any]]) -> list[str]:
    return ["Ad Soyadı"] + [str(d.get("day", "")) for d in days] + ["Çalıştığı", "Aylık", "Fazla"]


def _row_totals(
    row: dict[str, Any],
    days: list[dict[str, Any]],
    cells_map: dict[str, str],
) -> tuple[int, int, int]:
    worked = 0
    for d in days:
        code = cells_map.get(d.get("iso") or "", "") or ""
        if code == "8":
            worked += 8
        elif code == "16":
            worked += 16
        elif code == "24":
            worked += 24
        elif code == "Yİ":
            worked += 8
    ideal = int(row.get("ideal_hours") or 0)
    if row.get("role") == "lead":
        ideal = 0
    ot = max(0, worked - ideal) if row.get("role") != "lead" else 0
    if "worked_hours" in row:
        worked = int(row.get("worked_hours") or worked)
    if "overtime_hours" in row and row.get("role") != "lead":
        ot = int(row.get("overtime_hours") or ot)
    if "ideal_hours" in row and row.get("role") != "lead":
        ideal = int(row.get("ideal_hours") or ideal)
    return worked, ideal, ot


def _ensure_empty_lead_export_rows(
    days: list[dict[str, Any]],
    rows: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """İndirmede Gülten satırı her zaman en üstte ve tamamen boş."""
    empty_cells = {str(d.get("iso") or ""): "" for d in days if d.get("iso")}
    lead_row = {
        "name": LEAD_NURSE,
        "role": "lead",
        "cells": empty_cells,
        "worked_hours": 0,
        "ideal_hours": 0,
        "overtime_hours": 0,
    }
    others = [
        r
        for r in (rows or [])
        if (r.get("name") or "") != LEAD_NURSE and r.get("role") != "lead"
    ]
    return [lead_row, *others]


def _export_matrix(
    days: list[dict[str, Any]],
    rows: list[dict[str, Any]],
) -> tuple[list[str], list[list[Any]]]:
    headers = _export_headers(days)
    body: list[list[Any]] = []
    for row in _ensure_empty_lead_export_rows(days, rows):
        cells_map = row.get("cells") or {}
        worked, ideal, ot = _row_totals(row, days, cells_map)
        line: list[Any] = [row.get("name") or ""]
        for d in days:
            line.append(cells_map.get(d.get("iso") or "", "") or "")
        line.extend([worked, ideal, ot])
        body.append(line)
    return headers, body


def build_ayilma_csv_bytes(
    *,
    year: int,
    month: int,
    days: list[dict[str, Any]],
    rows: list[dict[str, Any]],
) -> bytes:
    """UTF-8 BOM + virgül ayırıcı; boş hücreler \"\" — bitişik virgül (,,) yok."""
    import csv
    from io import StringIO

    headers, body = _export_matrix(days, rows)
    sio = StringIO()
    sio.write("\ufeff")
    writer = csv.writer(
        sio,
        delimiter=",",
        lineterminator="\r\n",
        quoting=csv.QUOTE_ALL,
        quotechar='"',
    )
    writer.writerow([_month_title(month, year)])
    writer.writerow([])
    writer.writerow(headers)
    for line in body:
        writer.writerow(line)
    return sio.getvalue().encode("utf-8")


def build_ayilma_docx_bytes(
    *,
    year: int,
    month: int,
    days: list[dict[str, Any]],
    rows: list[dict[str, Any]],
) -> bytes:
    """Word .docx — yatay A4, sabit sütun; Windows Word / Android uyumlu."""
    from io import BytesIO

    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Mm, Pt

    def _set_nowrap(cell) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        for old in tc_pr.findall(qn("w:noWrap")):
            tc_pr.remove(old)
        tc_pr.append(OxmlElement("w:noWrap"))

    def _write_cell(cell, text: str, *, bold: bool = False, size: float = 8, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_nowrap(cell)

    def _fixed_layout(table) -> None:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

    headers, body = _export_matrix(days, rows)
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(0.7)
    section.right_margin = Cm(0.7)

    title = doc.add_heading(_month_title(month, year), level=1)
    for run in title.runs:
        run.font.size = Pt(14)

    table = doc.add_table(rows=1 + len(body), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _fixed_layout(table)

    name_w = Mm(26)
    total_w = Mm(13)
    n_days = len(days)
    usable = section.page_width - section.left_margin - section.right_margin
    day_budget = usable - name_w - (total_w * 3)
    day_w = day_budget // max(n_days, 1)
    day_w = max(day_w, Mm(4.8))

    col_widths = [name_w] + [day_w] * n_days + [total_w, total_w, total_w]
    for col_idx, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[col_idx].width = width

    wd_tr = ("Pzt", "Sal", "Çar", "Per", "Cum", "Cmt", "Paz")
    hdr = table.rows[0].cells
    _write_cell(hdr[0], "Ad Soyadı", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    for i, d in enumerate(days):
        wd = wd_tr[d.get("weekday", 0) % 7] if isinstance(d.get("weekday"), int) else ""
        label = f"{d.get('day', '')}\n{wd}" if wd else str(d.get("day", ""))
        _write_cell(hdr[1 + i], label, bold=True, size=7)
    for j, label in enumerate(("Çalıştığı", "Aylık", "Fazla")):
        _write_cell(hdr[1 + n_days + j], label, bold=True, size=7)

    for r_i, line in enumerate(body):
        cells = table.rows[r_i + 1].cells
        _write_cell(cells[0], str(line[0]), size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        for c_i, d in enumerate(days):
            val = line[1 + c_i]
            _write_cell(cells[1 + c_i], "" if val == "" else str(val), size=8)
        for j in range(3):
            _write_cell(cells[1 + n_days + j], str(line[-3 + j]), size=8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
